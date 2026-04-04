import pdfplumber
import fitz          
import pytesseract
from PIL import Image
from docx import Document
import io
import os

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
MIN_CHARS_THRESHOLD = 100


def extract_text(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return _extract_from_pdf_smart(file_path)
    elif ext in (".docx", ".doc"):
        return _extract_from_docx(file_path)
    else:
        raise ValueError(f"Format non supporté : {ext}")


def _extract_from_pdf_smart(file_path: str) -> str:
    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page_number, page in enumerate(pdf.pages):

            native_text = page.extract_text()

            if native_text and len(native_text.strip()) >= MIN_CHARS_THRESHOLD:
                text_parts.append(native_text)
                print(f"  Page {page_number + 1} : texte natif ({len(native_text)} chars)")

            else:
                print(f"  Page {page_number + 1} : OCR nécessaire...")
                ocr_text = _ocr_pdf_page(file_path, page_number)
                if ocr_text:
                    text_parts.append(ocr_text)
                    print(f"  Page {page_number + 1} : OCR OK ({len(ocr_text)} chars)")
                else:
                    print(f"  Page {page_number + 1} : OCR vide (image illisible ?)")
    return "\n".join(text_parts).strip()


def _ocr_pdf_page(file_path: str, page_number: int) -> str:
   
    doc = fitz.open(file_path)
    page = doc[page_number]

    mat = fitz.Matrix(2, 2)
    pix = page.get_pixmap(matrix=mat)

    img_data = pix.tobytes("png")
    img = Image.open(io.BytesIO(img_data))

    doc.close()

    text = pytesseract.image_to_string(
        img,
        lang="fra+eng",   
        config="--psm 6"  
    )
    return text.strip()


def _extract_from_docx(file_path: str) -> str:
    doc = Document(file_path)
    parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip() and cell.text not in parts:
                    parts.append(cell.text)
    for section in doc.sections:
        for p in section.header.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for p in section.footer.paragraphs:
            if p.text.strip():
                parts.append(p.text)
    return "\n".join(parts).strip()