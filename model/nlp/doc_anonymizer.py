# nlp/doc_anonymizer.py
import fitz
import re
import spacy
from langdetect import detect
from docx import Document
import os

nlp_fr = spacy.load("fr_core_news_md")
nlp_en = spacy.load("en_core_web_md")

SENSITIVE_PATTERNS = {  
    r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b':                  "[EMAIL]",
    r'(\+212|\+33|0)([ \-\.]?\d){9,10}':                                      "[TELEPHONE]",
    r'(\+\d{1,3}[\s\-]?)?\(?\d{3}\)?[\s\-]?\d{3}[\s\-]?\d{4}':              "[TELEPHONE]",
    r'\b(1[89]|[2-5]\d)\s*ans?\b':                                            "[AGE]",
    r'\b(1[89]|[2-5]\d)\s*years?\s*old\b':                                    "[AGE]",
    r'(né[e]?\s*le|born\s*on)\s*\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{2,4}':       "[DATE_NAISSANCE]",
    r'(https?://)?(www\.)?(linkedin\.com/in/|github\.com/)[A-Za-z0-9_\-/]+': "[PROFIL]",
}

SKILLS_KEYWORDS = [
    "python", "java", "javascript", "typescript", "c++", "c#", "php",
    "ruby", "swift", "kotlin", "scala", "go", "rust", "r",
    "react", "angular", "vue", "html", "css", "bootstrap", "tailwind",
    "spring boot", "flask", "django", "node.js", "express", "laravel",
    "mysql", "postgresql", "mongodb", "redis", "oracle", "sql server", "sql",
    "machine learning", "deep learning", "nlp", "tensorflow", "pytorch",
    "scikit-learn", "pandas", "numpy", "xgboost", "keras",
    "docker", "kubernetes", "aws", "azure", "gcp", "git", "jenkins",
    "nginx", "linux", "ci/cd",
]

DIPLOMAS_FR = ["master", "licence", "bac", "baccalauréat", "ingénieur",
               "doctorat", "bts", "dut", "mba", "phd", "mpsi", "cpge"]
DIPLOMAS_EN = ["bachelor", "master", "phd", "doctorate", "mba",
               "associate degree", "high school", "diploma", "degree",
               "engineering degree", "bsc", "msc", "ba", "ma", "computer science"]

EXPERIENCE_PATTERNS_FR = [
    r'(\d+)\s*an[s]?\s*d[\'e]expérience',
    r'(\d{4})\s*[-–]\s*(\d{4}|présent|aujourd\'hui|actuel)',
]
EXPERIENCE_PATTERNS_EN = [
    r'(\d+)\s*year[s]?\s*of\s*experience',
    r'(\d{4})\s*[-–]\s*(\d{4}|present|current|now)',
]


def create_anonymized_file(input_path: str, output_path: str) -> dict:
    ext = os.path.splitext(input_path)[1].lower()
    if ext == ".pdf":
        return _anonymize_pdf(input_path, output_path)
    elif ext in (".docx", ".doc"):
        return _anonymize_docx(input_path, output_path)
    else:
        raise ValueError(f"Format non supporté : '{ext}'. Utilisez .pdf ou .docx.")


def _anonymize_pdf(input_path: str, output_path: str) -> dict:
    doc = fitz.open(input_path)
    full_text = ""
    for page in doc:
        full_text += page.get_text()

    language  = _detect_language(full_text)
    nlp       = nlp_fr if language == "fr" else nlp_en
    spacy_doc = nlp(full_text)

    person_names = list(set([
        ent.text for ent in spacy_doc.ents
        if ent.label_ == "PER" and len(ent.text.strip()) > 2
    ]))
    organisations = list(set([
        ent.text for ent in spacy_doc.ents if ent.label_ == "ORG"
    ]))

    for page in doc:
        page_text = page.get_text()

        for name in person_names:
            for rect in page.search_for(name):
                page.add_redact_annot(rect, text="[NOM]", fontsize=9,
                                      fill=(1, 1, 1), text_color=(0.1, 0.1, 0.8))

        for pattern, label in SENSITIVE_PATTERNS.items():
            for match in re.finditer(pattern, page_text, re.IGNORECASE):
                for rect in page.search_for(match.group()):
                    page.add_redact_annot(rect, text=label, fontsize=9,
                                          fill=(1, 1, 1), text_color=(0.8, 0.1, 0.1))
        page.apply_redactions()

    doc.save(output_path)
    doc.close()

    fields = _extract_fields(full_text, language)
    fields["organisations"] = organisations
    return {"language": language, "fields": fields}


def _anonymize_docx(input_path: str, output_path: str) -> dict:
    doc = Document(input_path)
    full_text = _extract_full_docx_text(doc)

    language  = _detect_language(full_text)
    nlp       = nlp_fr if language == "fr" else nlp_en
    spacy_doc = nlp(full_text)

    person_names = list(set([
        ent.text for ent in spacy_doc.ents
        if ent.label_ == "PER" and len(ent.text.strip()) > 2
    ]))
    organisations = list(set([
        ent.text for ent in spacy_doc.ents if ent.label_ == "ORG"
    ]))

    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, person_names)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, person_names)

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            _replace_in_paragraph(paragraph, person_names)
        for paragraph in section.footer.paragraphs:
            _replace_in_paragraph(paragraph, person_names)

    doc.save(output_path)

    fields = _extract_fields(full_text, language)
    fields["organisations"] = organisations
    return {"language": language, "fields": fields}


def _extract_full_docx_text(doc: Document) -> str:
    parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    for section in doc.sections:
        for p in section.header.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for p in section.footer.paragraphs:
            if p.text.strip():
                parts.append(p.text)
    return "\n".join(parts)


def _replace_in_paragraph(paragraph, person_names: list):
    full_text = paragraph.text
    if not full_text.strip():
        return
    anonymized = full_text
    for name in person_names:
        anonymized = anonymized.replace(name, "[NOM]")
    for pattern, label in SENSITIVE_PATTERNS.items():
        anonymized = re.sub(pattern, label, anonymized, flags=re.IGNORECASE)
    if anonymized != full_text and paragraph.runs:
        paragraph.runs[0].text = anonymized
        for run in paragraph.runs[1:]:
            run.text = ""


def _detect_language(text: str) -> str:
    try:
        lang = detect(text[:500])
        return "fr" if lang == "fr" else "en"
    except Exception:
        return "fr"


def _extract_fields(text: str, language: str) -> dict:
    text_lower = text.lower()
    diplomas_list = DIPLOMAS_FR if language == "fr" else DIPLOMAS_EN
    exp_patterns  = EXPERIENCE_PATTERNS_FR if language == "fr" else EXPERIENCE_PATTERNS_EN

    found_experiences = []
    for pattern in exp_patterns:
        matches = re.findall(pattern, text_lower)
        if matches:
            found_experiences.extend([str(m) for m in matches])

    return {
        "competences": [s for s in SKILLS_KEYWORDS if s in text_lower],
        "diplomes":    [d for d in diplomas_list if d in text_lower],
        "experiences": found_experiences,
        "langue_cv":   language,
    }